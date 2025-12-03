from tkinter import messagebox, filedialog
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd

class docxGenerator:
    def __init__(self, template_path, cliente, numero_proposta, proposta_completa, data_label, servicos):
        self.template_path = template_path
        self.cliente = cliente
        self.numero_proposta = numero_proposta
        self.proposta_completa = proposta_completa
        self.data_label = data_label
        self.servicos = servicos

    # ==================== DOCX ====================
    def replace_text_keep_formatting(self, paragraph, placeholder, new_text):
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, new_text)

    def replace_placeholder_formatted(self, document, placeholder, new_text):
        for p in document.paragraphs:
            if placeholder in p.text:
                self.replace_text_keep_formatting(p, placeholder, new_text)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if placeholder in p.text:
                            self.replace_text_keep_formatting(p, placeholder, new_text)

    def gerar_docx(self):
        tpl = self.template_path.get()
        if not tpl or not os.path.isfile(tpl):
            messagebox.showwarning('Aviso', 'Selecione um modelo .docx válido')
            return
        if not self.cliente.get().strip():
            messagebox.showwarning('Aviso', 'Informe o nome do cliente')
            return
        if not self.numero_proposta.get().strip():
            messagebox.showwarning('Aviso', 'Informe o número da proposta')
            return

        try:
            doc = Document(tpl)
        except Exception as e:
            messagebox.showerror('Erro', f'Não foi possível abrir o modelo: {e}')
            return

        cliente_txt = self.cliente.get().upper()
        proposta_txt = self.proposta_completa.get()
        data_txt = self.data_label.get()
        self.replace_placeholder_formatted(doc, '{{NOME}}', cliente_txt)
        self.replace_placeholder_formatted(doc, '{{PROPOSTA}}', proposta_txt)
        self.replace_placeholder_formatted(doc, '{{DATA}}', data_txt)

        df = pd.DataFrame(self.servicos)
        if df.empty:
            messagebox.showwarning('Aviso', 'Adicione pelo menos um serviço')
            return
        valor_total = df['Total (R$)'].sum()

        tabela = None
        for table in doc.tables:
            try:
                if len(table.columns) >= 5:
                    tabela = table
                    break
            except Exception:
                continue

        if tabela:
            try:
                while len(tabela.rows) > 2:
                    tabela._tbl.remove(tabela.rows[-2]._tr)
                for i, row_data in enumerate(self.servicos, start=1):
                    last_row = tabela.rows[-1]._tr
                    new_row = tabela.add_row()._tr
                    tabela._tbl.remove(new_row)
                    last_row.addprevious(new_row)
                    row_cells = tabela.rows[-2].cells
                    larg = row_data['Largura'] if row_data['Largura'] not in [None, '', '0', 0] else 'X'
                    alt = row_data['Altura'] if row_data['Altura'] not in [None, '', '0', 0] else 'X'
                    dados = [str(i), row_data['Descrição'], str(larg), str(alt),
                            str(row_data['Quantidade']), f'R$ {row_data["Preço"]:.2f}', f'R$ {row_data["Total (R$)"]:.2f}']

                    for idx_col, cell in enumerate(row_cells):
                        cell.text = ''
                        p = cell.paragraphs[0]
                        p.text = dados[idx_col]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                total_row = tabela.rows[-1].cells
                total_row[-2].text = 'TOTAL'
                total_row[-1].text = f'R$ {valor_total:,.2f}'
                for c in total_row:
                    for p in c.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                messagebox.showwarning('Aviso', f'Erro ao preencher tabela: {e}')
                tabela = None

        if not tabela:
            doc.add_paragraph('SERVIÇOS:')
            for i, row_data in enumerate(self.servicos, start=1):
                doc.add_paragraph(f"{i} - {row_data['Descrição']} | LxA: {row_data['Largura']}x{row_data['Altura']} | Qtd: {row_data['Quantidade']} | R$ {row_data['Total (R$)']:,.2f}")
            doc.add_paragraph(f"TOTAL: R$ {valor_total:,.2f}")

        default_name = f"orcamento_{cliente_txt.replace(' ', '_')}_{proposta_txt}.docx"
        save_path = filedialog.asksaveasfilename(defaultextension=".docx", initialfile=default_name, filetypes=[('Word Document', '*.docx')])
        if not save_path:
            return
        try:
            doc.save(save_path)
            messagebox.showinfo('Sucesso', f'Orçamento gerado: {save_path}')
        except Exception as e:
            messagebox.showerror('Erro', f'Não foi possível salvar o arquivo: {e}')